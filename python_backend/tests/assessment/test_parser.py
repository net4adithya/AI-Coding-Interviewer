# tests/assessment/test_parser.py
import pytest
from io import BytesIO
from docx import Document

from app.assessment.parsers.docx_parser import DocxQuestionBankParser

def create_mock_docx() -> BytesIO:
    doc = Document()
    doc.add_heading("Question 1: Reverse Array", level=1)
    doc.add_paragraph("Topic: Arrays")
    doc.add_paragraph("Difficulty: EASY")
    doc.add_paragraph("Time: 15 mins")
    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph("Reverse a given array.")
    doc.add_heading("Starter Code", level=2)
    doc.add_paragraph("def reverse(arr):\n    pass")
    doc.add_heading("Test Cases", level=2)
    doc.add_paragraph("[1, 2, 3]")
    
    doc.add_heading("Question: Two Sum", level=1)
    doc.add_paragraph("Topic: Arrays")
    doc.add_paragraph("Difficulty: MEDIUM")
    doc.add_paragraph("Time: 20 mins")
    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph("Find two sum.")
    doc.add_heading("Test Cases", level=2)
    doc.add_paragraph("[1, 2, 3], 4")
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def test_docx_parser():
    buf = create_mock_docx()
    parser = DocxQuestionBankParser()
    result = parser.parse(buf)
    
    assert len(result.errors) == 0
    assert len(result.questions) == 2
    
    q1 = result.questions[0]
    assert q1.title == "Reverse Array"
    assert q1.topic == "Arrays"
    assert q1.difficulty == "EASY"
    assert q1.expected_time_minutes == 15
    assert q1.problem_statement == "Reverse a given array."
    assert "default" in q1.starter_code
    assert len(q1.test_cases) == 1
    assert q1.test_cases[0]["stdin"] == "[1, 2, 3]"

    q2 = result.questions[1]
    assert q2.title == "Two Sum"
    assert q2.topic == "Arrays"
    assert q2.difficulty == "MEDIUM"
    assert q2.expected_time_minutes == 20
    assert len(q2.test_cases) == 1
